import os
import sys

from docx import Document
from ssp import SSP
from ssp.control import Control

TEST_DIR = os.path.dirname(os.path.abspath(__file__))

class TestControl(object):

    def test_control_creation(self):
        #need to make a better test ssp for this, just a quick test...
        ssp = Document(TEST_DIR + r'/test_files/blank_templates/08282018/FedRAMP-SSP-High-Baseline-Template.docx')
        cis_table = ''
        implementation_table = ''
        next_table_is_implementation_details = False
        for table in ssp.tables:
            try:
                if table.cell(0,0).text == 'AC-1':
                    next_table_is_implementation_details = True
                    cis_table = table
                elif next_table_is_implementation_details:
                    implementation_table = table
                    test_control = Control(cis_table, implementation_table)
                    break
            except IndexError:
                pass

        assert test_control.number == 'AC-1'
        assert test_control.cis_table == cis_table
        assert test_control.implementation_table == implementation_table
        assert len(test_control.parts) == len(implementation_table.rows) -1


    def test_part(self):
        ssp = Document(TEST_DIR + r'/test_files/blank_templates/08282018/FedRAMP-SSP-High-Baseline-Template.docx')
        cis_table = ''
        implementation_table = ''
        next_table_is_implementation_details = False
        for table in ssp.tables:
            try:
                if table.cell(0,0).text == 'AC-2':
                    next_table_is_implementation_details = True
                    cis_table = table
                elif next_table_is_implementation_details:
                    implementation_table = table
                    test_control = Control(cis_table, implementation_table)
                    break
            except IndexError:
                pass

        assert test_control.part('d').text == ''

    def test_responsible_role(self):
        ssp = SSP(TEST_DIR + r'/test_files/blank_templates/08282018/FedRAMP-SSP-High-Baseline-Template.docx')
        assert ssp.control('AC-1').responsible_role == "My Role"

    def test_parameters(self):
        ssp = SSP(TEST_DIR + r'/test_files/blank_templates/08282018/FedRAMP-SSP-High-Baseline-Template.docx')
        assert ssp.control('AC-1').parameters == ['Param AC-1(a)','Param AC-1(b)(1)','Param AC-1(b)(2)']

    def test_implementation_status(self):
        ssp = SSP(TEST_DIR + r'/test_files/blank_templates/08282018/FedRAMP-SSP-High-Baseline-Template.docx')
        assert ssp.control('AC-1').implementation_status == ["Implemented"]

    def test_control_origination(self):
        ssp = SSP(TEST_DIR + r'/test_files/blank_templates/08282018/FedRAMP-SSP-High-Baseline-Template.docx')
        assert ssp.control('AC-1').control_origination == ["Service Provider Corporate"]
    