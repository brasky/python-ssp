from docx import Document
from ssp.control import Control

class SecurityPlan(object):
    """
    Not intended to be constructed directly. Use ssp.api.SSP() to open an SSP.
    """

    control_list = []
    control_list_to_table_index = {}

    def __init__(self, path):
        self.source = path
        self.document = Document(path)
        self.version = self.get_version()
        self.child_ssp = None
        if self.version == '08282018':
            self.child_ssp = SecurityPlan_08282018(self)
        else:
            raise ValueError('This template version is not compatible.')
        self.control_list = self.child_ssp.control_list
        self.control_list_to_table_index = self.child_ssp.control_list_to_table_index

    def __iter__(self):
        return iter(self.control_list)
    
    def control(self, control):
        """
        Takes a control as a string and returns the matching control object in the SSP.
        """
        try:
            return self.control_list[self.control_list_to_table_index[control.upper()]]
        except KeyError:
            raise KeyError('No control found with that name')

    def get_version(self):
        """
        Needs to take a Document file and return either a version number or raise an error if not an ssp.
        """
        # For some reason docx doesn't see anything before the table of contents in the list of tables or paragraphs, will need to parse document._element.xml
        # template_revision_table_index = self.table_name_to_table_index['Date']
        # template_revision_table = self.tables[template_revision_table_index]
        # try:
        #     date = template_revision_table.cell(7, 0).text
        #     if date == '8/28/2018':
        #         return '08282018'
        #     else:
        #         raise ValueError
        # except:
        #     raise ValueError('Error getting version, has the template revision history been altered? Possibly unsupported template')
        return '08282018'

class SecurityPlan_08282018(SecurityPlan):
    """
    Child class for the SSP version published 08/28/2018. 
    Takes an SSP object and returns a child SSP object specific to this template.
    """
    
    def __init__(self, ssp):
        self.document = ssp.document
        self.tables = ssp.document.tables
        self.create_control_index()

    def is_cis_table(self, table):
        try:
            if '-' in table.cell(0,0).text and len(table.cell(0,0).text) < 9:
                return True
            return False
        except:
            return False

    def create_control_index(self):
        """
        Creates control objects for each pair of CIS/Implementation tables 
        and adds them to control_list and the list index. 
        """
        implementation_table_next = False
        cis_table = ''
        for table in self.document.tables:
            try:
                if self.is_cis_table(table):
                    implementation_table_next = True
                    cis_table = table
                elif implementation_table_next:
                    new_control = Control(cis_table, table)
                    self.control_list_to_table_index[new_control.number] = len(self.control_list)
                    self.control_list.append(new_control)
                    implementation_table_next = False
            except Exception as e:
                print(e)
